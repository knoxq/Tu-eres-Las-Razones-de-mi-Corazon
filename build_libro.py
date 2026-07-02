import json, subprocess, sys, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))

def load_book():
    js = subprocess.run(
        ["node", "-e", "import('./src/data/capitulos.js').then(m=>process.stdout.write(JSON.stringify({prologo:m.prologo,capitulos:m.capitulos,epilogo:m.epilogo})))"],
        cwd=HERE, capture_output=True, text=True, check=True,
    )
    return json.loads(js.stdout)

def add_section(doc, section, heading_prefix=""):
    # Título
    title = heading_prefix + section['titulo'] if heading_prefix else section['titulo']
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x6B, 0x21, 0x3F)

    # Contenido (sin subheading "Contenido" para prólogo/epílogo)
    if section.get("contenido"):
        for para in section["contenido"].split("\n"):
            if para.strip() == "":
                continue
            p = doc.add_paragraph(para.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.size = Pt(12)

    # Carta / Al lector (solo capítulos)
    if section.get("carta"):
        doc.add_heading("Al lector", level=2)
        for para in section["carta"].split("\n"):
            if para.strip() == "":
                continue
            p = doc.add_paragraph(para.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.size = Pt(12)
                run.italic = True

    # Separador
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.add_run("❦").font.size = Pt(16)

def add_chapter(doc, cap):
    # Resumen (en cursiva)
    if cap.get("resumen"):
        p = doc.add_paragraph()
        r = p.add_run(cap["resumen"])
        r.italic = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_section(doc, cap, heading_prefix=f"Capítulo {cap['id']}: ")

def main():
    book = load_book()
    doc = Document()

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(12)

    # Portada
    title = doc.add_heading("Tú Eres las Razones de Mi Corazón", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("— una historia contada en tercera persona hasta que ya no se pudo —")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.italic = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x6B, 0x21, 0x3F)
    doc.add_page_break()

    # Prólogo
    if book.get("prologo"):
        add_section(doc, book["prologo"])

    # Capítulos
    capitulos = book.get("capitulos", [])
    for cap in capitulos:
        add_chapter(doc, cap)

    # Epílogo
    if book.get("epilogo"):
        add_section(doc, book["epilogo"])

    out = os.path.join(HERE, "libro.docx")
    doc.save(out)
    print(f"Guardado: {out} ({len(capitulos)} capítulos + prólogo + epílogo)")

if __name__ == "__main__":
    main()
